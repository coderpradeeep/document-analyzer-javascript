import io
import tempfile
from PyPDF2 import PdfReader
import docx

def extract_text_from_pdf(file):
    """
    Extract text from a PDF file.
    
    Args:
        file: The uploaded PDF file object from Streamlit
        
    Returns:
        str: The extracted text from the PDF
    """
    try:
        # Create a PDF reader object
        pdf = PdfReader(io.BytesIO(file.getvalue()))
        
        # Initialize an empty string to store the text
        text = ""
        
        # Iterate through each page and extract text
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:  # Only add if text was extracted
                text += page_text + "\n\n"
        
        return text
    except Exception as e:
        # Add more specific error handling for different PDF errors
        raise Exception(f"Error extracting text from PDF: {str(e)}")

def extract_text_from_docx(file):
    """
    Extract text from a DOCX file.
    
    Args:
        file: The uploaded DOCX file object from Streamlit
        
    Returns:
        str: The extracted text from the DOCX file
    """
    try:
        # Save the uploaded file to a temporary file
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
            temp_file.write(file.getvalue())
            temp_file_path = temp_file.name
        
        # Open the document with python-docx
        doc = docx.Document(temp_file_path)
        
        # Extract text from paragraphs
        full_text = []
        for para in doc.paragraphs:
            if para.text:  # Only add non-empty paragraphs
                full_text.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():  # Only add non-empty cells
                        full_text.append(cell.text.strip())
        
        # Join all text with newlines
        return '\n'.join(full_text)
    except Exception as e:
        raise Exception(f"Error extracting text from DOCX: {str(e)}")
